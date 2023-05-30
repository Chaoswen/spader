import os
import requests
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm

# 添加段落样式
def set_paragraph_style(paragraph, font_name, font_size, is_bold, alignment):
    if len(paragraph.runs) > 0:
        run = paragraph.runs[0]  # 获取段落的第一个运行对象
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = is_bold
        paragraph.alignment = alignment

# 定义文件路径
file_path = "C:/download/2020/2020年x月.docx"  # 替换为你的文档路径

# 检查文件是否存在
if not os.path.exists(file_path):
    # 创建一个新的 Word 文档
    document = Document()
else:
    # 打开现有的 Word 文档
    document = Document(file_path)

# 存放链接的 URL 数组

urls = [
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AgtkZGRil9TN9BejGb5Jlw96cDhreZc2lw%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=1&__EVENTVALIDATION=%2FwEWBQLMx6mwBQKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrAdPYbpAMDGGTwwtZaeATG0Y9OM8x&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=11&lblPaperID=1234",
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AgFkZGRzb0xiqVtNjXWknRqzDddRONxjzg%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=2&__EVENTVALIDATION=%2FwEWBQLG7rmHAwKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrAWs1ppBTr3hoFP1EJWAYyZuU9Aya&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=1&lblPaperID=1234",
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AgJkZGSc0rXygZBfouLJ7gJkTias5Zf0OQ%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=3&__EVENTVALIDATION=%2FwEWBQLkwo3qBAKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrAUekWPotYN83TaLeRQfJ%2BezZMeCt&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=2&lblPaperID=1234",
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AgNkZGT0Lr%2F05HCJHGJtTRdT2fuMg6YaPw%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=4&__EVENTVALIDATION=%2FwEWBQKyosesDgKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrAaXeGco2yje6SGr1mjh61Uii9XIc&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=3&lblPaperID=1234",
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AgRkZGTgnQfOzVcFZ8hl6Bes3s4gN17Fpg%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=5&__EVENTVALIDATION=%2FwEWBQLew9O1AgKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrARwpuUvsRG2Lm41V9PpmV3u1q%2FAp&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=4&lblPaperID=1234",
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AgVkZGQiWQu8fAIH7sKrRm%2B4V2scbRYzRw%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=6&__EVENTVALIDATION=%2FwEWBQKBjPWiBgKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrASosLLg5NeWP5cao%2Fu5Qbc6GqeDI&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=5&lblPaperID=1234",
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AgZkZGTt7Kjff2Ln7f6z9LciA%2FtJ5zIfdg%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=7&__EVENTVALIDATION=%2FwEWBQKTlIWDDwKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrAV8bgW87byZ6%2FPnlN%2BZLTo%2BibDRe&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=6&lblPaperID=1234",
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AgdkZGQzHwWk8u53szKfIYumJ0mJwFLUlQ%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=8&__EVENTVALIDATION=%2FwEWBQLh2pWAAgKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrAX2L2MjPqG%2Bl6vhXdcjFcMtn0Tys&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=7&lblPaperID=1234",
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AghkZGTsikbXCK%2B7VeGri3WGWiXCVsGr7g%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=9&__EVENTVALIDATION=%2FwEWBQKNtLeaDQKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrAQXfEwBIkYbttziRmC8DxkOH8zbm&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=8&lblPaperID=1234",
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AglkZGT0pS%2B%2Fwn%2FilVS0x96xVHiCdrpYZg%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=10&__EVENTVALIDATION=%2FwEWBQKKnYCiCAKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrAfdMTcpFuUBcfsyxuR%2BRmaSQZM2C&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=9&lblPaperID=1234",
    "http://fullsearch.cnepaper.com/FullSearch.aspx?__VIEWSTATE=%2FwEPDwULLTE4NTgxMDgzMjQPZBYCAgEPZBYCAgMPDxYEHgtSZWNvcmRjb3VudAJnHhBDdXJyZW50UGFnZUluZGV4AgpkZGT21skkUxVeq7cq%2FFsdoNFB02X0NA%3D%3D&__VIEWSTATEGENERATOR=4F4D99FC&__EVENTTARGET=AspNetPager1&__EVENTARGUMENT=11&__EVENTVALIDATION=%2FwEWBQL33pC3BwKMwc%2FlAQK2hea3DQL4p5WKCgLY0YCrAR3sjtyX96ZLJiRSttoec9DHC17n&search_text=%E8%84%B1%E8%B4%AB%2C%E6%89%B6%E8%B4%AB&Txt_SiteStart=2020-03-01&Txt_SiteEnd=2020-03-31&AspNetPager1_input=10&lblPaperID=1234"
]

# 爬取并解析页面
for url in tqdm(urls, desc="Progress"):
    response = requests.get(url)
    content = response.content
    # 使用Beautiful Soup解析网页内容
    soup = BeautifulSoup(response.content, "html.parser")

    

    # 处理解析后的网页内容
    # 这里可以根据网页的结构和你的需求来提取相关信息
    # 以下示例提取网页中的所有链接
    links = soup.select("a:not([href^='javascript:__doPostBack()'])")
    previous_date = None  # 保存前一个日期

    for link in links:
        link_url = link.get("href")
        if link_url is None:
            continue


        # 检查链接是否以"javascript:"开头，如果是，则跳过处理
        if link_url.startswith("javascript:"):
            continue

        # 发送HTTP GET请求获取链接对应页面的内容
        try:
            link_response = requests.get(link_url)
            link_content = link_response.content

            # 使用Beautiful Soup解析链接页面的内容
            link_soup = BeautifulSoup(link_content, "html.parser")

            # 提取指定<span>标签的文本内容
            span_element = link_soup.find("span", class_="default")
            if span_element:
                text = span_element.text

                # 判断日期是否与前一个日期相同，如果不同则打印日期
                if text != previous_date:
                    #print(text)
                    previous_date = text
                    # 输出到文档
                    paragraph = document.add_paragraph(text)
                    # 设置日期样式
                    set_paragraph_style(paragraph, "SimSun", 10.5, True, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

            # 提取文章标题
            title = link_soup.find("p", class_="BSHARE_TEXT")
            if title is not None:
                paragraph = document.add_paragraph(title.text)
                # 设置标题样式
                set_paragraph_style(paragraph, "SimHei", 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER)

            # 提取文章内容
            content = link_soup.find("founder-content")
            if content is not None:
                # 替换连续两个空格为换行符
                content_text = re.sub(r'\s{2}', '\n', content.text)
                # 添加换行符后添加到文档中
                paragraphs = content_text.split("\n")
                for para in paragraphs:
                    if para.strip():
                        paragraph = document.add_paragraph()
                        # 设置内容样式
                        set_paragraph_style(paragraph, "SimSun", 10.5, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
                        # 添加首行缩进两个字符
                        run = paragraph.add_run("\u3000" * 2 + para)  


        except requests.exceptions.RequestException as e:
            print(f"无法获取链接内容: {link_url}")
            print(e)

    # 保存文档
    document.save(file_path)

